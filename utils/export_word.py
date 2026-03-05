"""
Word document export for CyberPresales AI
Generates a professional proposal document from analysis results
"""
import os
import tempfile
from datetime import datetime

def generate_word_doc(session_state):
    """Generate a Word document from all available analysis results"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy
    except ImportError:
        import subprocess
        subprocess.run(["pip", "install", "python-docx", "--break-system-packages", "-q"])
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Color palette
    NAVY = RGBColor(0x1e, 0x40, 0xaf)
    DARK = RGBColor(0x1e, 0x29, 0x3b)
    GRAY = RGBColor(0x64, 0x74, 0x8b)
    WHITE = RGBColor(0xff, 0xff, 0xff)

    def add_heading(text, level=1, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = color or NAVY
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = color or NAVY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = color or DARK
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(16)
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '93C5FD')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_label(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.bold = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        return p

    # ── Cover Page ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    run = p.add_run("CYBERSECURITY PROPOSAL")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run("Presales Intelligence Report")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)

    file_name = getattr(session_state, 'file_name', None) or session_state.get('file_name', 'RFP Document')
    p = doc.add_paragraph()
    run = p.add_run(f"Document: {file_name}")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}  ·  CyberPresales AI v3.0  ·  by Yash Mehrotra")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(40)

    add_divider()
    doc.add_page_break()

    # ── Helper to add section from session state ──────────────────────────────
    def add_section(title, content, section_num):
        add_heading(f"{section_num}. {title}", level=1)
        add_divider()
        if content:
            # Split by ## headings for better formatting
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith('## '):
                    add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip('*'))
                    run.bold = True
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = DARK
                else:
                    add_body(line)
        else:
            p = doc.add_paragraph()
            run = p.add_run("Analysis not yet generated. Run this module in the app to populate this section.")
            run.font.size = Pt(10)
            run.font.color.rgb = GRAY
            run.italic = True

    # Get content from session state
    def get(key):
        if hasattr(session_state, key):
            return getattr(session_state, key)
        if hasattr(session_state, '__getitem__'):
            return session_state.get(key)
        return None

    section_num = 1

    if get('exec_summary'):
        add_section("Executive Summary", get('exec_summary'), section_num)
        section_num += 1
        doc.add_page_break()

    if get('customer_brief'):
        add_section("Customer Intelligence Brief", get('customer_brief'), section_num)
        section_num += 1
        doc.add_page_break()

    if get('pain_analysis'):
        add_section("Pain Point & Requirements Analysis", get('pain_analysis'), section_num)
        section_num += 1
        doc.add_page_break()

    if get('solution_rec'):
        add_section("Solution Recommendation", get('solution_rec'), section_num)
        section_num += 1
        doc.add_page_break()

    if get('competitive'):
        add_section("Competitive Landscape", get('competitive'), section_num)
        section_num += 1
        doc.add_page_break()

    if get('product_map'):
        add_section("Product & Vendor Mapping", get('product_map'), section_num)
        section_num += 1

    # Footer note
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("— End of Report —")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)

    p = doc.add_paragraph()
    run = p.add_run("Generated by CyberPresales AI · Built by Yash Mehrotra · Powered by Groq LLaMA 3.3")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    output_path = os.path.join(tempfile.gettempdir(), "CyberPresales_Proposal.docx")
    doc.save(output_path)
    return output_path
