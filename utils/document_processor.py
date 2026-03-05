import io

def extract_text_from_file(uploaded_file):
    """Extract text from PDF or DOCX uploaded file."""
    try:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == 'pdf':
            return extract_from_pdf(uploaded_file)
        elif file_type == 'docx':
            return extract_from_docx(uploaded_file)
        else:
            return None
    except Exception as e:
        print(f"Error extracting text: {e}")
        return None


def extract_from_pdf(uploaded_file):
    """Extract text from PDF file."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip() if text.strip() else None
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return None


def extract_from_docx(uploaded_file):
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(uploaded_file.read()))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip() if text.strip() else None
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return None


def extract_multiple_files(uploaded_files):
    """
    Extract and merge text from multiple uploaded files.
    Returns: merged text, list of (filename, word_count, status)
    """
    merged_text = ""
    doc_summaries = []

    for uploaded_file in uploaded_files:
        try:
            text = extract_text_from_file(uploaded_file)
            if text:
                word_count = len(text.split())
                # Add document separator so AI knows which doc is which
                merged_text += f"""
\n\n{'='*60}
DOCUMENT: {uploaded_file.name}
{'='*60}\n
{text}
"""
                doc_summaries.append({
                    "name": uploaded_file.name,
                    "words": word_count,
                    "status": "success",
                    "size": f"{uploaded_file.size // 1024} KB"
                })
            else:
                doc_summaries.append({
                    "name": uploaded_file.name,
                    "words": 0,
                    "status": "failed",
                    "size": f"{uploaded_file.size // 1024} KB"
                })
        except Exception as e:
            doc_summaries.append({
                "name": uploaded_file.name,
                "words": 0,
                "status": "error",
                "size": "unknown"
            })

    return merged_text.strip(), doc_summaries

